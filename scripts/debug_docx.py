"""
debug_docx.py
Run this on a generated document to see exactly what text
is in each paragraph and table cell. This tells us whether
the placeholder "{{ ttd_pejabat }}" is actually present in the
file and what form it takes in the XML.

Usage:
    python3 scripts/debug_docx.py public/cached_result/your_file.docx
"""

import sys
import zipfile
import re
from pathlib import Path

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 debug_docx.py <path_to_docx>")
        sys.exit(1)

    path = sys.argv[1]

    if not Path(path).exists():
        print(f"ERROR: File not found: {path}")
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"Inspecting: {path}")
    print(f"{'='*60}\n")

    # ── 1. Raw XML inspection — find ALL occurrences of ttd/qr ──
    print("── RAW XML SEARCH ──────────────────────────────────────────")
    with zipfile.ZipFile(path, 'r') as z:
        for name in z.namelist():
            if name.endswith('.xml'):
                try:
                    content = z.read(name).decode('utf-8', errors='replace')
                    # Search for anything that looks like our placeholders
                    if any(kw in content for kw in ['ttd', 'qr_code', 'nama_pejabat', 'jabatan', 'tgl_ttd', '{{']):
                        print(f"\n  Found in: {name}")
                        # Show lines containing the keywords
                        for i, line in enumerate(content.splitlines()):
                            if any(kw in line for kw in ['ttd', 'qr_code', 'nama_pejabat', 'jabatan', 'tgl_ttd', '{{']):
                                # Strip XML tags to show just the text content
                                clean = re.sub(r'<[^>]+>', '', line).strip()
                                if clean:
                                    print(f"    Line {i}: {clean[:120]}")
                                else:
                                    print(f"    Line {i} (raw): {line[:120]}")
                except Exception as e:
                    print(f"  Could not read {name}: {e}")

    print()

    # ── 2. python-docx paragraph inspection ──────────────────────
    print("── PYTHON-DOCX PARAGRAPH TEXT ──────────────────────────────")
    try:
        from docx import Document
        doc = Document(path)

        print("\n  Body paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            full = "".join(r.text for r in para.runs)
            if full.strip():
                print(f"    [{i}] '{full}'")
                # Also show individual runs if paragraph looks suspicious
                if '{{' in full or 'ttd' in full or 'qr' in full:
                    print(f"         ^ FOUND PLACEHOLDER — runs breakdown:")
                    for j, run in enumerate(para.runs):
                        print(f"           run[{j}]: '{run.text}'")

        print("\n  Table cells:")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        full = "".join(r.text for r in para.runs)
                        if full.strip():
                            print(f"    table[{t_idx}] row[{r_idx}] col[{c_idx}] para[{p_idx}]: '{full}'")
                            if '{{' in full or 'ttd' in full or 'qr' in full:
                                print(f"         ^ FOUND PLACEHOLDER — runs breakdown:")
                                for j, run in enumerate(para.runs):
                                    print(f"           run[{j}]: '{run.text}'")

    except ImportError:
        print("  python-docx not installed. Run: venv/bin/pip install python-docx")
    except Exception as e:
        print(f"  Error reading with python-docx: {e}")

    print()

    # ── 3. Check for split runs in raw XML ───────────────────────
    print("── CHECKING FOR SPLIT RUNS (Word XML fragmentation) ────────")
    with zipfile.ZipFile(path, 'r') as z:
        if 'word/document.xml' in z.namelist():
            content = z.read('word/document.xml').decode('utf-8', errors='replace')
            # Find all <w:t> tag contents
            wt_values = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', content)
            print(f"\n  All <w:t> text values (non-empty):")
            for v in wt_values:
                if v.strip():
                    print(f"    '{v}'")

            # Check if {{ or }} appear split across tags
            print(f"\n  Checking if {{ or }} are split across <w:t> tags:")
            # Concatenate all w:t values and look for our placeholders
            joined = ''.join(wt_values)
            placeholders = ['ttd_pejabat', 'qr_code', 'nama_pejabat', 'jabatan_pejabat', 'tgl_ttd']
            for ph in placeholders:
                if ph in joined:
                    print(f"    FOUND '{ph}' in concatenated text")
                else:
                    print(f"    NOT FOUND '{ph}' in concatenated text")

    print(f"\n{'='*60}")
    print("Debug complete.")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()